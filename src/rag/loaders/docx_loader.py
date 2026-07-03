"""DOCX 格式加载器（python-docx）"""
from __future__ import annotations

import logging
import re
from typing import TYPE_CHECKING, List

from src.rag.data_sources import FileInfo
from src.rag.loaders.base import BaseLoader, register_loader
from src.rag.outline import OutlineTree, extract_docx_headings

if TYPE_CHECKING:
    from langchain_core.documents import Document as _Doc

logger = logging.getLogger(__name__)

# DOCX heading style names (varies by locale, English names are most common)
_HEADING_STYLES = {
    "heading 1", "heading 2", "heading 3", "heading 4", "heading 5", "heading 6",
    "标题 1", "标题 2", "标题 3", "标题 4", "标题 5", "标题 6",  # Chinese
}


@register_loader(".docx")
class DocxLoader(BaseLoader):
    """加载 DOCX 文件，提取文字 + 结构信息 + 大纲

    处理流程：
        1. 遍历段落，识别 Heading 样式提取标题列表
        2. 构建大纲树
        3. 按章节边界拆分文档（每个章节从 Heading 到下一个同级或更高级 Heading）
    """

    def load(self, info: FileInfo, base_meta: dict) -> List["_Doc"]:
        from src.rag.loader import (
            _filter_noise_paragraphs,
            normalize_text,
        )
        from langchain_core.documents import Document

        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning(
                "python-docx not installed. Install with: pip install python-docx"
            )
            return []

        try:
            doc = DocxDocument(str(info.path))
        except Exception as e:
            logger.warning("Failed to open DOCX %s: %s", info.path, e)
            return []

        # 提取文档属性
        props = doc.core_properties
        if props.author:
            base_meta["author"] = props.author
        if props.title:
            base_meta["title"] = props.title
        if props.created:
            base_meta["created_time"] = props.created.isoformat()
        if props.modified:
            base_meta["modified_time"] = props.modified.isoformat()

        # 1. 提取标题（从 Heading 样式）
        headings = extract_docx_headings(doc.paragraphs)

        # 2. 构建章节列表
        # 每个段落标记为 (level, text, style_name)
        paragraph_entries: List[tuple] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = (para.style.name or "").lower()

            if style_name in _HEADING_STYLES:
                # 提取层级数字
                match = re.search(r"(\d+)", style_name)
                level = int(match.group(1)) if match else 1
                level = min(level, 6)
                paragraph_entries.append(("heading", level, text))
            else:
                paragraph_entries.append(("body", 0, text))

        # 3. 按章节分组
        chapters = self._group_by_chapters(paragraph_entries, headings)

        # 4. 如果没有标题，整篇作为一个文档
        if not chapters:
            full_text = "\n\n".join(
                t for _, _, t in paragraph_entries
            )
            full_text = normalize_text(full_text)
            full_text = _filter_noise_paragraphs(full_text)
            if not full_text.strip():
                return []
            return [Document(page_content=full_text, metadata={**base_meta, "source_file": info.name})]

        # 5. 构建大纲树并拆分
        outline_tree = OutlineTree()
        outline_tree.build(headings)
        source_name = info.name

        # 对于 DOCX，我们手动构建 Document 列表（因为段落结构不同于文本块）
        docs: List["_Doc"] = []
        for chapter_idx, (chapter_heading, chapter_level, chapter_text) in enumerate(chapters):
            chapter_text = normalize_text(chapter_text)
            chapter_text = _filter_noise_paragraphs(chapter_text)
            if not chapter_text.strip():
                continue

            # 查找 chapter_path
            flat = outline_tree.flatten()
            chapter_path = ""
            for path, ht, lvl in flat:
                if ht == chapter_heading and lvl == chapter_level:
                    chapter_path = path
                    break
            if not chapter_path:
                chapter_path = chapter_heading

            meta = {
                **base_meta,
                "source_file": source_name,
                "chapter_path": chapter_path,
                "heading_level": chapter_level,
                "heading_text": chapter_heading,
            }
            if flat:
                from src.rag.config import settings
                store_json = getattr(settings, "outline_store_full_json", False)
                if store_json:
                    meta["outline"] = json.dumps(
                        outline_tree.root.to_dict() if outline_tree.root else {},
                        ensure_ascii=False,
                    )

            docs.append(Document(page_content=chapter_text, metadata=meta))

        return docs

    def _group_by_chapters(
        self,
        paragraphs: List[tuple],
        headings: List[tuple],
    ) -> List[tuple]:
        """将段落按章节分组

        返回 [(heading_text, heading_level, body_text), ...]
        每个章节包含标题和其后的正文段落。
        """
        if not headings:
            return []

        # 构建章节起始位置索引
        heading_positions = {i for i, (kind, level, _) in enumerate(paragraphs) if kind == "heading"}

        chapters: List[tuple] = []
        heading_items = [(i, l, t) for i, (kind, l, t) in enumerate(paragraphs) if kind == "heading"]

        for idx, (pos, level, text) in enumerate(heading_items):
            # 找到下一个 heading 的位置
            next_pos = heading_items[idx + 1][0] if idx + 1 < len(heading_items) else len(paragraphs)

            # 收集该 heading 后的 body 段落
            body_lines: List[str] = []
            for j in range(pos + 1, next_pos):
                kind, _, t = paragraphs[j]
                if kind == "body" and t.strip():
                    body_lines.append(t.strip())

            body_text = "\n\n".join(body_lines)
            chapters.append((text, level, body_text))

        return chapters
